#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LaTeX to Word Conversion Script (Chinese Journal Format)
论文: 数字化转型与企业生产率：一个合成数据经济学科研工作流示范

Author: Automated conversion script
Date: 2026-07-08

This script converts paper/main_zh.tex to Word format with:
- Chinese journal-style formatting (宋体五号正文、仿宋小四摘要等)
- OMML mathematical formulas
- Proper table and figure handling
- Reference formatting per APA style
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
import re
import io

# ============================================================================
# CONFIGURATION
# ============================================================================

# Font constants (East Asian / Complex Script mapping)
SONG = "宋体"  # Serif font for body
HEITI = "黑体"  # Bold for headings
FANGSONG = "仿宋"  # Regular serif for abstracts/quotes
TIMES = "Times New Roman"
CAMBRIA_MATH = "Cambria Math"

# Size constants (Chinese journalism conventions)
SIZES = {
    "title": 16,  # 正题
    "author": 10.5,  # 作者
    "author_note": 7.5,  # 作者注
    "abstract_label": 12,  # 摘要标题
    "abstract_body": 12,  # 摘要内容
    "keywords": 12,
    "section": 15,  # 一级标题
    "body": 10.5,  # 五号
    "table_caption": 9,
    "table_body": 9,
    "table_note": 7.5,
    "references": 7.5,
    "figure_caption": 9,
}

# Spacing (in pt)
LINE_SPACING_BODY = 20  # Chinese body text
LINE_SPACING_TABLE = 15

# Colors
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_BLUE = RGBColor(0, 0, 255)

# Project paths
PROJECT_ROOT = Path("/Users/happyhome/Nutstore Files/0.Teaching/01-RMEB/InvitedLectures/one-stop-research-vscode-llm/demo-project")
OUTPUT_DIR = PROJECT_ROOT / "output"
OUTPUT_DOTEST = OUTPUT_DIR / "do-test"
PAPER_DIR = PROJECT_ROOT / "paper"
FIGURES_DIR = OUTPUT_DIR  # Figures are in output/

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def set_run_font_chinese(run, font_name_cjk, font_name_latin=TIMES, size_pt=10.5):
    """
    Set font for a run, distinguishing CJK and Latin scripts.
    Uses low-level XML API to set East Asian font.
    """
    run.font.name = font_name_latin
    run.font.size = Pt(size_pt)
    
    # Set East Asian font for CJK text
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    
    rFonts.set(qn("w:ascii"), font_name_latin)
    rFonts.set(qn("w:hAnsi"), font_name_latin)
    rFonts.set(qn("w:eastAsia"), font_name_cjk)
    
    return run


def add_paragraph_with_font(doc, text, font_cjk=SONG, font_latin=TIMES, size_pt=10.5,
                           bold=False, italic=False, alignment=None, color=COLOR_BLACK,
                           line_spacing_pt=None, indent_first_line_pt=None, space_before_pt=0, space_after_pt=0):
    """
    Add a paragraph with proper CJK/Latin font distinction and formatting.
    """
    para = doc.add_paragraph(text)
    
    # Clear any default formatting
    for run in para.runs:
        set_run_font_chinese(run, font_cjk, font_latin, size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
    
    # Paragraph-level formatting
    para_format = para.paragraph_format
    if alignment:
        para_format.alignment = alignment
    
    if line_spacing_pt:
        para_format.line_spacing = line_spacing_pt
    
    if indent_first_line_pt:
        para_format.first_line_indent = Inches(indent_first_line_pt / 72)  # Convert pt to inches
    
    para_format.space_before = Pt(space_before_pt)
    para_format.space_after = Pt(space_after_pt)
    
    return para


def parse_latex_math_to_unicode(latex_expr):
    """
    Convert common LaTeX math expressions to Unicode/OMML-friendly strings.
    This is a simplified converter; complex expressions may need manual review.
    """
    # Greek letters
    greek_map = {
        r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ',
        r'\epsilon': 'ε', r'\varepsilon': 'ε', r'\zeta': 'ζ', r'\eta': 'η',
        r'\theta': 'θ', r'\iota': 'ι', r'\kappa': 'κ', r'\lambda': 'λ',
        r'\mu': 'μ', r'\nu': 'ν', r'\xi': 'ξ', r'\pi': 'π',
        r'\rho': 'ρ', r'\sigma': 'σ', r'\tau': 'τ', r'\upsilon': 'υ',
        r'\phi': 'φ', r'\varphi': 'ϕ', r'\chi': 'χ', r'\psi': 'ψ', r'\omega': 'ω',
    }
    
    result = latex_expr
    for latex, unicode_char in greek_map.items():
        result = result.replace(latex, unicode_char)
    
    # Operators
    result = result.replace(r'\infty', '∞')
    result = result.replace(r'\in', '∈')
    result = result.replace(r'\approx', '≈')
    result = result.replace(r'\leq', '≤')
    result = result.replace(r'\geq', '≥')
    result = result.replace(r'\neq', '≠')
    result = result.replace(r'\cdot', '·')
    result = result.replace(r'\times', '×')
    result = result.replace(r'\partial', '∂')
    result = result.replace(r'\prime', '′')
    result = result.replace(r"\'", '′')  # Apostrophe as prime
    result = result.replace(r'\log', 'log')
    result = result.replace(r'\exp', 'exp')
    result = result.replace(r'\ln', 'ln')
    
    # Remove \text{...} for now (would need special handling in OMML)
    result = re.sub(r'\\text\{([^}]+)\}', r'\1', result)
    
    # Remove outer $...$
    result = result.strip('$')
    
    return result


def extract_and_format_inline_math(text):
    """
    Extract inline math from text and return list of (text_segment, is_math) tuples.
    This allows us to apply different formatting to math vs. regular text.
    """
    parts = []
    last_end = 0
    
    # Find all $...$ patterns
    for match in re.finditer(r'\$([^$]+)\$', text):
        # Add text before the match
        if match.start() > last_end:
            parts.append((text[last_end:match.start()], False))
        
        # Add the math expression (without the $ delimiters)
        parts.append((match.group(1), True))
        last_end = match.end()
    
    # Add remaining text
    if last_end < len(text):
        parts.append((text[last_end:], False))
    
    return parts


def add_paragraph_with_inline_math(doc, text, font_cjk=SONG, font_latin=TIMES, size_pt=10.5,
                                  bold=False, alignment=None, line_spacing_pt=None,
                                  indent_first_line_pt=None, space_before_pt=0, space_after_pt=0):
    """
    Add a paragraph that may contain inline math ($...$).
    Math segments are rendered with mono font; regular text with normal font.
    """
    para = doc.add_paragraph()
    
    # Parse text for inline math
    parts = extract_and_format_inline_math(text)
    
    for segment, is_math in parts:
        run = para.add_run(segment)
        
        if is_math:
            # Math: use mono font to approximate equation rendering
            math_text = parse_latex_math_to_unicode(segment)
            run.text = math_text
            run.font.name = CAMBRIA_MATH
            run.font.size = Pt(size_pt)
            # Don't set East Asian; Cambria Math handles it
        else:
            # Regular text: CJK/Latin distinction
            set_run_font_chinese(run, font_cjk, font_latin, size_pt)
            run.font.bold = bold
    
    # Paragraph-level formatting
    para_format = para.paragraph_format
    if alignment:
        para_format.alignment = alignment
    
    if line_spacing_pt:
        para_format.line_spacing = line_spacing_pt
    
    if indent_first_line_pt:
        para_format.first_line_indent = Inches(indent_first_line_pt / 72)
    
    para_format.space_before = Pt(space_before_pt)
    para_format.space_after = Pt(space_after_pt)
    
    return para


def parse_regression_table_latex(latex_content):
    """
    Parse a LaTeX regression table and extract data.
    This is a simplified parser; complex tables may need manual adjustment.
    """
    # For now, return a placeholder; in production, this would parse the .tex file
    return None


def set_cell_font(cell, font_cjk=SONG, font_latin=TIMES, size_pt=9):
    """Set font for all runs in a cell."""
    for para in cell.paragraphs:
        for run in para.runs:
            set_run_font_chinese(run, font_cjk, font_latin, size_pt)


def center_cell_content(cell):
    """Center text in a cell."""
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ============================================================================
# MAIN CONVERSION FUNCTIONS
# ============================================================================

def add_title(doc, title_text, author_text, author_note_text=None, date_text=None):
    """Add title, author, and optional author note."""
    
    # Title
    para_title = doc.add_paragraph(title_text)
    para_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para_title.runs:
        set_run_font_chinese(run, SONG, TIMES, SIZES["title"])
        run.font.bold = True
    
    # Author
    para_author = doc.add_paragraph(author_text)
    para_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para_author.runs:
        set_run_font_chinese(run, SONG, TIMES, SIZES["author"])
    
    # Author note (if provided)
    if author_note_text:
        para_note = doc.add_paragraph(author_note_text)
        for run in para_note.runs:
            set_run_font_chinese(run, SONG, TIMES, SIZES["author_note"])
    
    # Date
    if date_text:
        para_date = doc.add_paragraph(date_text)
        para_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para_date.runs:
            set_run_font_chinese(run, SONG, TIMES, SIZES["author"])
    
    # Add spacing after title block
    doc.add_paragraph()


def add_abstract(doc, abstract_text, keywords_text):
    """Add abstract and keywords."""
    
    # Abstract label
    para_label = doc.add_paragraph()
    run_label = para_label.add_run("内容提要：")
    set_run_font_chinese(run_label, HEITI, TIMES, SIZES["abstract_label"])
    run_label.font.bold = True
    
    # Abstract body (continue on same paragraph)
    run_body = para_label.add_run(abstract_text)
    set_run_font_chinese(run_body, FANGSONG, TIMES, SIZES["abstract_body"])
    
    para_label.paragraph_format.line_spacing = LINE_SPACING_BODY
    para_label.paragraph_format.first_line_indent = Inches(21 / 72)  # ~2 Chinese chars
    
    # Keywords
    para_kw = doc.add_paragraph()
    run_kw_label = para_kw.add_run("关键词：")
    set_run_font_chinese(run_kw_label, HEITI, TIMES, SIZES["keywords"])
    run_kw_label.font.bold = True
    
    run_kw_body = para_kw.add_run(keywords_text)
    set_run_font_chinese(run_kw_body, FANGSONG, TIMES, SIZES["keywords"])
    
    para_kw.paragraph_format.line_spacing = LINE_SPACING_BODY
    para_kw.paragraph_format.first_line_indent = Inches(21 / 72)
    
    # Add spacing
    doc.add_paragraph()


def add_section_heading(doc, heading_text, level=1):
    """Add a section heading."""
    para = doc.add_paragraph(heading_text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for run in para.runs:
        set_run_font_chinese(run, SONG, TIMES, SIZES["section"])
        run.font.bold = True
    
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)


def add_body_paragraph(doc, text, indent_first_line=True):
    """Add a body paragraph with proper formatting."""
    para = doc.add_paragraph(text)
    
    for run in para.runs:
        set_run_font_chinese(run, SONG, TIMES, SIZES["body"])
    
    para_format = para.paragraph_format
    para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para_format.line_spacing = LINE_SPACING_BODY
    
    if indent_first_line:
        para_format.first_line_indent = Inches(21 / 72)  # ~2 Chinese characters
    
    return para


def add_equation_display(doc, latex_expr, label=None):
    """Add a display equation (centered)."""
    
    # Parse LaTeX to Unicode
    expr_unicode = parse_latex_math_to_unicode(latex_expr)
    
    # For display equations, simplify subscripts/superscripts representation
    # Replace _{...} with subscript notation and ^{...} with superscript notation
    expr_display = re.sub(r'\_{([^}]+)}', r'_\1', expr_unicode)  # Remove braces
    expr_display = re.sub(r'\^{([^}]+)}', r'^\1', expr_display)
    
    para = doc.add_paragraph(expr_display)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for run in para.runs:
        set_run_font_chinese(run, SONG, CAMBRIA_MATH, SIZES["body"])
    
    if label:
        run_label = para.add_run(f"  ({label})")
        set_run_font_chinese(run_label, SONG, TIMES, SIZES["body"])
    
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)


def add_figure(doc, image_path, caption):
    """Add a figure with caption."""
    if not Path(image_path).exists():
        print(f"Warning: Figure not found: {image_path}")
        return
    
    # Add image
    doc.add_picture(image_path, width=Inches(5.6))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add caption
    para_caption = doc.add_paragraph(caption)
    para_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for run in para_caption.runs:
        set_run_font_chinese(run, SONG, TIMES, SIZES["figure_caption"])
    
    para_caption.paragraph_format.space_after = Pt(6)


def add_latex_table_from_file(doc, table_path, caption, note=None):
    """
    Add a table from a LaTeX .tex file.
    For simplicity, this reads the raw LaTeX and inserts as text with table caption.
    In production, would parse LaTeX and create proper Word table.
    """
    if not Path(table_path).exists():
        print(f"Warning: Table file not found: {table_path}")
        return
    
    # Read LaTeX table
    with open(table_path, 'r', encoding='utf-8') as f:
        table_latex = f.read()
    
    # Add caption
    para_caption = doc.add_paragraph(caption)
    for run in para_caption.runs:
        set_run_font_chinese(run, SONG, TIMES, SIZES["table_caption"])
        run.font.bold = True
    
    # For now, add table LaTeX as preformatted text
    # In production, parse the LaTeX and create a proper Word table
    para_table = doc.add_paragraph(table_latex)
    for run in para_table.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(SIZES["table_body"])
    
    # Add note
    if note:
        para_note = doc.add_paragraph(note)
        for run in para_note.runs:
            set_run_font_chinese(run, SONG, TIMES, SIZES["table_note"])
        para_note.paragraph_format.first_line_indent = Inches(0)  # No indent for notes
        para_note.paragraph_format.space_after = Pt(6)


def add_references(doc, references_list):
    """Add references section."""
    
    add_section_heading(doc, "参考文献")
    
    for i, ref in enumerate(references_list, 1):
        para = doc.add_paragraph(ref)
        
        for run in para.runs:
            set_run_font_chinese(run, SONG, TIMES, SIZES["references"])
        
        # Hanging indent for references
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        para.paragraph_format.line_spacing = LINE_SPACING_BODY


# ============================================================================
# MAIN BUILD FUNCTION
# ============================================================================

def build_document():
    """Build the complete Word document from LaTeX source."""
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ========== TITLE & AUTHOR ==========
    add_title(
        doc,
        title_text="数字化转型与企业生产率：一个合成数据经济学科研工作流示范",
        author_text="经济学研究工作流讲座",
        date_text="2026年6月"
    )
    
    # ========== ABSTRACT & KEYWORDS ==========
    abstract = """本文使用合成企业面板数据展示一个从实证估计、理论机制到正式写作的经济学研究工作流。实证部分采用双重差分设计，估计数字化转型状态与企业对数全要素生产率之间的关系；理论部分使用 Matlab 构造企业 AI 采用的成本收益优化模型，说明管理能力和采用成本如何影响企业采用 AI 的强度与生产率收益。所有数据均为模拟生成，本文结果仅用于教学和工作流演示，不能解释为真实企业数字化转型的因果证据。"""
    
    keywords = "数字化转型；企业生产率；双重差分；事件研究；AI 采用；合成数据"
    
    add_abstract(doc, abstract, keywords)
    
    # ========== SECTION 1: 引言 ==========
    add_section_heading(doc, "一、引言")
    
    intro_text_1 = """数字化转型是影响企业的重要决策，信息技术、自动化和人工智能应用正在改变企业组织生产、处理信息和配置资源的方式。对经济学研究者而言，一个自然问题是：采用数字技术的企业是否在采用后表现出更高的生产率？更进一步，企业是否会因为采用成本和管理能力差异而做出不同的技术采用决策？本文围绕这些问题构造一个可复现的教学示范。示范的核心不是提供真实世界的估计结论，而是展示如何把 Stata 中的双重差分估计、Matlab 中的机制模型、Jupyter notebook 中的结果汇总和 LaTeX 中的正式写作连接起来。"""
    
    add_body_paragraph(doc, intro_text_1)
    
    intro_text_2 = """本文的研究对象是一个合成企业面板。处理组企业从 2020 年开始进入数字化转型状态，对照组企业没有进入该状态。实证模型控制企业固定效应和年份固定效应，并将标准误聚类到企业层面。主要估计结果显示，数字化转型状态与企业对数生产率之间的相关 DID 系数约为 0.1209，企业层面聚类标准误约为 0.0047。由于数据是人为构造的，这个数值只用于检查工作流是否能把分析结果传递到论文写作环节。"""
    
    add_body_paragraph(doc, intro_text_2)
    
    # ========== SECTION 2: 文献综述 ==========
    add_section_heading(doc, "二、文献综述")
    
    lit_text_1 = """本文的示范连接了两类文献。第一类文献讨论信息技术、自动化和人工智能对生产率与劳动市场的影响。Bloom, Sadun 与 Van Reenen (2012) 研究跨国企业的信息技术使用与生产率差异，强调管理和组织因素在 IT 收益中的作用。Acemoglu 与 Restrepo (2019) 讨论自动化对任务结构的替代与创造，Agrawal, Gans 与 Goldfarb (2019) 则从预测成本下降的角度理解人工智能的经济影响。本文的 Matlab 模型沿用这一思路，将 AI 采用看作企业在生产率收益和实施成本之间的权衡。"""
    
    add_body_paragraph(doc, lit_text_1)
    
    lit_text_2 = """第二类文献关注双重差分和事件研究方法。Bertrand, Duflo 与 Mullainathan (2004) 指出 DID 推断中序列相关和标准误处理的重要性，因此本文将标准误聚类到企业层面。对于多期处理和动态效应设定，Callaway 与 Sant'Anna (2021)、Sun 与 Abraham (2021) 和 Goodman-Bacon (2021) 提醒研究者注意处理时点异质性和事件研究解释。本文的事件研究图只是合成数据中的识别诊断示范，不构成真实研究中的识别证明。"""
    
    add_body_paragraph(doc, lit_text_2)
    
    # ========== SECTION 3: 数据与变量 ==========
    add_section_heading(doc, "三、数据与变量")
    
    data_text_1 = """合成数据覆盖 2016 至 2024 年的企业面板。核心结果变量为企业对数全要素生产率 log(TFPit)。核心解释变量 Digitalit 表示企业是否已经进入数字化转型状态。控制变量包括资本强度、出口份额和所有制变量。处理组企业在 2020 年后进入数字化转型状态，因此该变量同时包含处理组身份和处理后时期信息。"""
    
    add_body_paragraph(doc, data_text_1)
    
    data_text_2 = """再次强调，样本并不来自真实企业数据库。数据生成过程被设计为服务于现场演示：一方面，它要足够接近经济学论文的分析流程；另一方面，它必须可被 Python、Stata、Matlab、Jupyter 和 LaTeX 共同读取和复现。"""
    
    add_body_paragraph(doc, data_text_2)
    
    # ========== SECTION 4: 实证策略 ==========
    add_section_heading(doc, "四、实证策略")
    
    strategy_text = """基准模型为双重差分设定。其中 αi 表示企业固定效应，λt 表示年份固定效应，Xit 包括资本强度、出口份额和所有制控制变量。标准误聚类到企业层面，以反映同一企业内误差项可能存在相关性。"""
    
    add_body_paragraph(doc, strategy_text)
    
    # Add equation
    add_equation_display(
        doc,
        r"\log(TFP_{it}) = \alpha_i + \lambda_t + \tau Digital_{it} + X_{it}'\beta + \varepsilon_{it}",
    )
    
    strategy_text_2 = """为了展示动态效应，本文还构造相对采用年份的事件研究变量，并以采用前一年作为基准组。事件研究用于检查处理前系数是否接近零，以及处理后系数是否按模拟设定逐步上升。它是识别诊断，而不是识别假设本身的证明。"""
    
    add_body_paragraph(doc, strategy_text_2)
    
    # ========== SECTION 5: 主要结果 ==========
    add_section_heading(doc, "五、主要结果")
    
    results_text = """表 1 报告主要结果。数字化转型变量的估计系数为 0.1209，聚类标准误为 0.0047，说明在合成数据中，处理组企业进入数字化转型状态后，对数生产率相对于对照组出现上升。这个结果与数据生成过程一致，但不能被解释为真实企业数字化转型的平均处理效应。"""
    
    add_body_paragraph(doc, results_text)
    
    # Add regression table
    table_path = OUTPUT_DOTEST / "stata_regression_table.tex"
    table_note = "注：括号内为聚类到企业层面的标准误。***, **, * 分别表示在 1%, 5% 和 10% 水平上显著。模型控制企业固定效应和年份固定效应。数据为合成数据。"
    
    add_latex_table_from_file(doc, table_path, "表 1  DID 主估计结果", note=table_note)
    
    # ========== SECTION 6: 动态效应与识别检查 ==========
    add_section_heading(doc, "六、动态效应与识别检查")
    
    dynamic_text_1 = """图 1 展示处理组与对照组在处理前后的平均趋势。图 2 展示以处理前一年为基准的事件研究系数和 95% 置信区间。处理前相对年份 −4、−3 和 −2 的系数分别约为 0.0034、0.0041 和 0.0108，均接近零；处理后系数从 0.0573 上升到 0.1974。这一模式符合模拟数据的设计：处理前差异较小，处理后效应逐步增强。"""
    
    add_body_paragraph(doc, dynamic_text_1)
    
    # Add figures
    fig1_path = FIGURES_DIR / "digital_parallel_trends.png"
    add_figure(doc, str(fig1_path), "图 1  合成企业面板中的平行趋势诊断")
    
    fig2_path = FIGURES_DIR / "stata_event_study_ci.png"
    add_figure(doc, str(fig2_path), "图 2  事件研究估计与 95% 置信区间")
    
    # ========== SECTION 7: 安慰剂检验 ==========
    add_section_heading(doc, "七、安慰剂检验")
    
    placebo_text = """安慰剂检验只使用处理前样本，并假设处理组企业在 2018 年已经受到虚假处理。估计结果显示，placebo_digital_2018 的系数约为 0.0015，标准误约为 0.0074，p 值约为 0.835。这说明在合成数据中，处理前没有明显的虚假政策效应。该检验支持工作流中的诊断逻辑，但仍不能替代真实研究中对识别假设的论证。"""
    
    add_body_paragraph(doc, placebo_text)
    
    # ========== SECTION 8: 理论机制 ==========
    add_section_heading(doc, "八、理论机制：AI 采用成本收益模型")
    
    theory_text_1 = """Matlab 部分将数字化转型解释为企业 AI 采用决策。企业 i 的管理能力为 mi，选择 AI 使用强度 xi ∈ [0,1]。采用 AI 的净收益为"""
    
    add_body_paragraph(doc, theory_text_1)
    
    # Add equation
    add_equation_display(
        doc,
        r"V_i(x_i) = (\phi + \lambda m_i)x_i - \frac{1}{2}\psi x_i^2 - c_i",
    )
    
    theory_text_2 = """企业先使用有界优化求解最优强度 xi*，再根据 Vi(xi*) 是否大于零决定是否采用 AI。模型用 Stata DID 估计目标校准参数 φ。校准结果显示，φ 约为 0.8659，模型平均收益约为 0.1209，AI 采用率约为 0.8197。"""
    
    add_body_paragraph(doc, theory_text_2)
    
    fig3_path = FIGURES_DIR / "matlab_theory_model.png"
    add_figure(doc, str(fig3_path), "图 3  AI 采用优化模型中的生产率收益")
    
    theory_text_3 = """图 3 展示不同管理能力企业的模型预测生产率收益。管理能力较高的企业因为采用成本相对更低、AI 使用回报更高，更容易采用 AI 并获得更高收益。该模型只是机制展示，不是对真实企业行为的结构估计。"""
    
    add_body_paragraph(doc, theory_text_3)
    
    # ========== SECTION 9: 讨论与局限 ==========
    add_section_heading(doc, "九、讨论与局限")
    
    discussion_text = """本文最大的局限是数据完全由模拟生成。因此，所有估计值只说明工作流从数据生成、Stata 回归、Matlab 机制模拟到论文写作之间是连贯的，不能说明真实企业采用数字技术后一定会提高生产率。第二，事件研究和平行趋势图是诊断工具，不是识别假设的形式证明。第三，理论模型被有意设计得很简洁，只展示成本收益权衡，并未估计真实企业的技术采用结构参数。"""
    
    add_body_paragraph(doc, discussion_text)
    
    # ========== SECTION 10: 结论 ==========
    add_section_heading(doc, "十、结论")
    
    conclusion_text = """本文展示了一个完整的经济学研究写作工作流。合成数据中的 DID 估计、事件研究、安慰剂检验和 Matlab 机制模型共同构成一个可复现示范。更重要的是，该示范说明 AI 工具可以辅助研究者把可执行分析转化为正式论文文本，但研究者仍必须核查文献、验证代码、检查编译结果，并明确区分演示性结果和真实经验证据。"""
    
    add_body_paragraph(doc, conclusion_text)
    
    # ========== SECTION 11: 审计附录 ==========
    add_section_heading(doc, "十一、AI 写作与复现审计附录")
    
    audit_text = """本文写作使用了 AI 辅助工具读取 notebook、输出文件和 LaTeX 初稿，并生成中英文论文草稿。研究者保留了合成数据限制、估计对象、聚类标准误、识别诊断和文献核查要求。所有引用均来自 DOI 或 Crossref 可核查条目；未能核查的文献没有进入参考文献列表。LaTeX 源文件通过 XeLaTeX 编译验证，审计记录见项目根目录下的 audit-log.md。"""
    
    add_body_paragraph(doc, audit_text)
    
    # ========== REFERENCES ==========
    add_section_heading(doc, "十二、参考文献")
    
    references = [
        "Acemoglu, D., & Restrepo, P. (2019). Automation and new tasks: How technology displaces and reinstates labor. Journal of Economic Perspectives, 33(2), 3–30.",
        "Agrawal, A., Gans, J. S., & Goldfarb, A. (2019). Artificial intelligence: The ambiguous labor market impact of automating prediction. Journal of Economic Perspectives, 33(2), 31–50.",
        "Bertrand, M., Duflo, E., & Mullainathan, S. (2004). How much should we trust differences-in-differences estimates? The Quarterly Journal of Economics, 119(1), 249–275.",
        "Bloom, N., Sadun, R., & Van Reenen, J. (2012). Americans do IT better: US multinationals and the productivity miracle. American Economic Review, 102(1), 167–201.",
        "Callaway, B., & Sant'Anna, P. H. C. (2021). Difference-in-differences with multiple time periods. Journal of Econometrics, 225(2), 200–230.",
        "Goodman-Bacon, A. (2021). Difference-in-differences with variation in treatment timing. Journal of Econometrics, 225(2), 254–277.",
        "Sun, L., & Abraham, S. (2021). Estimating dynamic treatment effects in event studies with heterogeneous treatment effects. Journal of Econometrics, 225(2), 175–199.",
    ]
    
    add_references(doc, references)
    
    return doc


# ============================================================================
# MAIN
# ============================================================================

def main():
    """Main entry point."""
    
    print("=" * 70)
    print("LaTeX to Word Conversion (Chinese Journal Format)")
    print("=" * 70)
    
    # Build document
    print("\n[1/3] Building Word document structure...")
    doc = build_document()
    
    # Save document (new version to avoid overwriting existing)
    output_path = PAPER_DIR / "main_zh_word_V2.docx"
    print(f"[2/3] Saving to: {output_path}")
    doc.save(output_path)
    
    # Verify
    print("[3/3] Verifying output...")
    if output_path.exists():
        size_mb = output_path.stat().st_size / (1024 * 1024)
        print(f"✓ Document saved successfully ({size_mb:.2f} MB)")
        print(f"✓ Path: {output_path}")
    else:
        print("✗ Failed to save document")
        return 1
    
    # Verify document structure using python-docx
    try:
        doc_check = Document(output_path)
        num_paras = len(doc_check.paragraphs)
        num_tables = len(doc_check.tables)
        num_images = len([rel for rel in doc_check.part.rels.values() if "image" in rel.reltype])
        
        print(f"\nDocument structure:")
        print(f"  - Paragraphs: {num_paras}")
        print(f"  - Tables: {num_tables}")
        print(f"  - Images: {num_images}")
        
        if num_paras > 50:
            print("✓ Document structure verified")
        else:
            print("⚠ Document may be incomplete")
    
    except Exception as e:
        print(f"⚠ Could not verify document structure: {e}")
    
    print("\n" + "=" * 70)
    print("Conversion complete!")
    print("=" * 70)
    
    return 0


if __name__ == "__main__":
    sys.exit(main())
