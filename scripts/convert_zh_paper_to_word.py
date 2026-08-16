from __future__ import annotations

import re
from pathlib import Path
from typing import Any, cast

from docx import Document as create_document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output"
PAPER = ROOT / "paper"
DOCX_PATH = PAPER / "main_zh_word.docx"

SONG = "宋体"
FANGSONG = "仿宋"
HEITI = "黑体"
TIMES = "Times New Roman"

# ── Math: OMML namespace ───────────────────────────────────────────
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

# ── Math: Greek letter lookup ──────────────────────────────────────
GREEK: dict[str, str] = {
    "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ",
    "epsilon": "ε", "varepsilon": "ε", "zeta": "ζ", "eta": "η",
    "theta": "θ", "iota": "ι", "kappa": "κ", "lambda": "λ",
    "mu": "μ", "nu": "ν", "xi": "ξ", "pi": "π",
    "rho": "ρ", "sigma": "σ", "tau": "τ", "upsilon": "υ",
    "phi": "φ", "varphi": "ϕ", "chi": "χ", "psi": "ψ",
    "omega": "ω",
    "Alpha": "Α", "Beta": "Β", "Gamma": "Γ", "Delta": "Δ",
    "Theta": "Θ", "Lambda": "Λ", "Xi": "Ξ", "Pi": "Π",
    "Sigma": "Σ", "Upsilon": "Υ", "Phi": "Φ", "Psi": "Ψ",
    "Omega": "Ω",
}

# ── Math: operator names that render upright ───────────────────────
OPERATORS = {"log", "ln", "exp", "sin", "cos", "tan", "lim", "max", "min", "det"}

# ── Math: special symbol translations ──────────────────────────────
SYMBOLS: dict[str, str] = {
    "cdot": "·", "times": "×", "div": "÷", "pm": "±",
    "leq": "≤", "geq": "≥", "neq": "≠", "approx": "≈",
    "equiv": "≡", "sim": "∼", "propto": "∝",
    "infty": "∞", "partial": "∂", "nabla": "∇",
    "rightarrow": "→", "leftarrow": "←", "Rightarrow": "⇒",
    "in": "∈", "notin": "∉", "subset": "⊂", "subseteq": "⊆",
    "cup": "∪", "cap": "∩", "forall": "∀", "exists": "∃",
    "angle": "∠", "parallel": "∥", "perp": "⊥",
    "prime": "′", "ast": "∗",
}


# ═══════════════════════════════════════════════════════════════════
# OMML construction helpers
# ═══════════════════════════════════════════════════════════════════

def _m(tag: str, *children, **attrs) -> OxmlElement:
    """Create an OMML element. tag is the local part (without 'm:' prefix)."""
    tag = tag.split(":")[-1]  # strip any existing m: prefix
    elm = OxmlElement(f"m:{tag}")
    for k, v in attrs.items():
        elm.set(k, v)
    for child in children:
        if isinstance(child, str):
            child = _m_run(child)
        elm.append(child)
    return elm


def _m_run(text: str) -> OxmlElement:
    """Create <m:r><m:t xml:space="preserve">text</m:t></m:r>."""
    r = OxmlElement("m:r")
    t = OxmlElement("m:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    return r


def _wrap_omath(*children) -> OxmlElement:
    """Wrap children in <m:oMath>."""
    omath = OxmlElement("m:oMath")
    for child in children:
        omath.append(child)
    return omath


def _wrap_omathpara(*children) -> OxmlElement:
    """Wrap children in <m:oMathPara>."""
    op = OxmlElement("m:oMathPara")
    for child in children:
        op.append(child)
    return op


# ═══════════════════════════════════════════════════════════════════
# LaTeX math → OMML parser (simple iterative scanner)
# ═══════════════════════════════════════════════════════════════════

def _match_braces(s: str, start: int) -> int:
    """Find the matching '}' for a '{' at position start-1. Returns position of '}'."""
    depth = 1
    i = start
    while i < len(s) and depth > 0:
        if s[i] == '{':
            depth += 1
        elif s[i] == '}':
            depth -= 1
        i += 1
    if depth != 0:
        raise ValueError(f"Unmatched brace in: {s}")
    return i - 1  # position of matching '}'


def _parse_latex_arg(s: str, start: int) -> tuple[str, int]:
    """Parse a LaTeX argument at start. Returns (content, new_pos).
    If s[start] == '{', content is inside braces. Otherwise single char."""
    if start >= len(s):
        return "", start
    if s[start] == '{':
        end = _match_braces(s, start + 1)
        return s[start + 1:end], end + 1
    return s[start], start + 1


def latex_to_omml(latex: str) -> OxmlElement:
    """Convert a LaTeX math string to an <m:oMath> element.
    
    Uses a simple iterative scan with manual subscript/superscript attachment.
    """
    return _build_omath(_scan_latex(latex))


def _scan_latex(s: str) -> list[OxmlElement]:
    """Scan LaTeX math string and return a flat list of OMML elements.
    
    Subscripts/superscripts are returned as marker elements (m:submark, m:supmark)
    that _build_omath will attach to their preceding base.
    """
    result: list[OxmlElement] = []
    i = 0
    n = len(s)
    
    while i < n:
        ch = s[i]
        
        # ── Subscript _ ──────────────────────────────────────────
        if ch == '_':
            i += 1
            arg_text, i = _parse_latex_arg(s, i)
            sub_children = _scan_latex(arg_text)
            marker = OxmlElement("m:submark")
            for c in sub_children:
                marker.append(c)
            result.append(marker)
            continue
        
        # ── Superscript ^ ────────────────────────────────────────
        if ch == '^':
            i += 1
            arg_text, i = _parse_latex_arg(s, i)
            sup_children = _scan_latex(arg_text)
            marker = OxmlElement("m:supmark")
            for c in sup_children:
                marker.append(c)
            result.append(marker)
            continue
        
        # ── LaTeX command \... ───────────────────────────────────
        if ch == '\\':
            i += 1
            # Read command name
            cmd_start = i
            while i < n and s[i].isalpha():
                i += 1
            cmd = s[cmd_start:i]
            
            if cmd in GREEK:
                result.append(_m_run(GREEK[cmd]))
            elif cmd in SYMBOLS:
                result.append(_m_run(SYMBOLS[cmd]))
            elif cmd in OPERATORS:
                # \log, \exp, \ln → m:func
                fname = _m("fName", _m_run(cmd))
                # Parse argument
                if i < n and s[i] in ('{', '('):
                    if s[i] == '(':
                        # consume until matching )
                        j = i + 1
                        depth = 1
                        while j < n and depth > 0:
                            if s[j] == '(': depth += 1
                            elif s[j] == ')': depth -= 1
                            j += 1
                        arg_text = s[i+1:j-1]
                        i = j
                    else:
                        arg_text, i = _parse_latex_arg(s, i)
                    arg_children = _build_omath_inline(_scan_latex(arg_text))
                    func_el = _m("func", fname, _m("e", *arg_children))
                else:
                    func_el = _m("func", fname, _m("e"))
                result.append(func_el)
            elif cmd == 'frac':
                num_text, i = _parse_latex_arg(s, i)
                den_text, i = _parse_latex_arg(s, i)
                num_kids = _scan_latex(num_text)
                den_kids = _scan_latex(den_text)
                result.append(_m("f", _m("num", *num_kids), _m("den", *den_kids)))
            elif cmd == 'sqrt':
                arg_text, i = _parse_latex_arg(s, i)
                kids = _scan_latex(arg_text)
                result.append(_m("rad", _m("e", *kids)))
            elif cmd in ('left', 'right'):
                # Skip \left( and \right) — just consume the delimiter
                if i < n:
                    i += 1
            elif cmd in ('begin', 'end', 'label', 'nonumber'):
                if i < n and s[i] == '{':
                    _, i = _parse_latex_arg(s, i)
            elif cmd == 'text' or cmd == 'texttt':
                arg_text, i = _parse_latex_arg(s, i)
                result.append(_m_run(arg_text))
            else:
                result.append(_m_run(cmd))
            continue
        
        # ── Special characters ───────────────────────────────────
        if ch == '\'':
            result.append(_m_run('′'))  # U+2032 prime
            i += 1
            continue
        
        if ch in ('{', '}'):
            i += 1  # skip stray braces
            continue
        
        if ch in ('(', ')', '[', ']', '=', '+', '-', ',', '|', ':', ' '):
            result.append(_m_run(ch))
            i += 1
            continue
        
        # ── Plain text run ───────────────────────────────────────
        j = i
        while j < n and s[j] not in ('_', '^', '\\', '\'', '{', '}',
                                       '(', ')', '[', ']', '=', '+', '-',
                                       ',', '|', ':', ' '):
            j += 1
        if j > i:
            result.append(_m_run(s[i:j]))
            i = j
            continue
        
        i += 1  # fallback
    
    return result


def _build_omath_inline(elements: list[OxmlElement]) -> list[OxmlElement]:
    """Like _build_omath but returns element list for nesting."""
    result: list[OxmlElement] = []
    
    for el in elements:
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        
        if tag == 'submark':
            if not result:
                continue
            base = result.pop()
            ssub = OxmlElement("m:sSub")
            ssub.append(_m("e", base))
            ssub.append(_m("sub", *list(el)))
            result.append(ssub)
        elif tag == 'supmark':
            if not result:
                continue
            base = result.pop()
            ssup = OxmlElement("m:sSup")
            ssup.append(_m("e", base))
            ssup.append(_m("sup", *list(el)))
            result.append(ssup)
        else:
            result.append(el)
    
    # Combine adjacent sSub + sSup → sSubSup
    merged: list[OxmlElement] = []
    i = 0
    while i < len(result):
        el = result[i]
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        
        if tag == 'sSub' and i + 1 < len(result):
            next_tag = result[i+1].tag.split('}')[-1] if '}' in result[i+1].tag else result[i+1].tag
            if next_tag == 'sSup':
                sub_el = result[i]
                sup_el = result[i+1]
                base = sub_el.find(qn('m:e'))
                sub = sub_el.find(qn('m:sub'))
                sup = sup_el.find(qn('m:sup'))
                ssubsup = OxmlElement("m:sSubSup")
                if base is not None:
                    ssubsup.append(base)
                if sub is not None:
                    ssubsup.append(sub)
                if sup is not None:
                    ssubsup.append(sup)
                merged.append(ssubsup)
                i += 2
                continue
        
        merged.append(el)
        i += 1
    
    return merged


def _build_omath(elements: list[OxmlElement]) -> OxmlElement:
    """Post-process: attach submark/supmark markers to preceding elements."""
    omath = OxmlElement("m:oMath")
    processed = _build_omath_inline(elements)
    for el in processed:
        omath.append(el)
    return omath


def add_inline_math(paragraph, latex: str, size: float = 10.5) -> None:
    """Insert an inline OMML equation into a paragraph."""
    omath = latex_to_omml(latex)
    run = paragraph.add_run()
    run.font.size = Pt(size)
    run._element.append(omath)


def add_display_equation(doc: DocxDocument, latex: str) -> None:
    """Add a centered display equation paragraph."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)

    omath = latex_to_omml(latex)
    ompara = _wrap_omathpara(_wrap_omath(*list(omath)))
    # ompara is the full <m:oMathPara> containing the equation
    paragraph._element.append(ompara)


def add_body_paragraph_with_math(doc: DocxDocument, text: str) -> None:
    """Add a body paragraph, rendering $...$ as inline OMML equations."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_format(paragraph)

    # Split on $...$ boundaries
    parts = re.split(r"(\$[^$]+\$)", text)
    for part in parts:
        if part.startswith("$") and part.endswith("$"):
            math_latex = part[1:-1]
            add_inline_math(paragraph, math_latex, size=10.5)
        else:
            add_mixed_run(paragraph, part, east_asia=SONG, size=10.5)


def add_abstract_paragraph_with_math(doc: DocxDocument, head_text: str, head_east_asia: str, head_size: float, body_text: str, body_east_asia: str, body_size: float) -> None:
    """Add an abstract/keywords paragraph with label in bold + body with inline math."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = Pt(20)

    # Label (bold)
    head = paragraph.add_run(head_text)
    set_run_font(head, east_asia=head_east_asia, size=head_size, bold=True)

    # Body with math
    parts = re.split(r"(\$[^$]+\$)", body_text)
    for part in parts:
        if part.startswith("$") and part.endswith("$"):
            math_latex = part[1:-1]
            add_inline_math(paragraph, math_latex, size=body_size)
        else:
            add_mixed_run(paragraph, part, east_asia=body_east_asia, size=body_size)


def set_run_font(run, east_asia: str = SONG, size: float = 10.5, bold: bool = False, italic: bool = False) -> None:
    run.font.name = TIMES
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_paragraph_format(paragraph, first_line: bool = True) -> None:
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(21)


def add_mixed_run(paragraph, text: str, east_asia: str = FANGSONG, size: float = 10.5, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=east_asia, size=size, bold=bold)


def add_body_paragraph(doc: DocxDocument, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_format(paragraph)
    add_mixed_run(paragraph, text, east_asia=SONG, size=10.5)


def add_section_heading(doc: DocxDocument, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=SONG, size=15, bold=True)


def add_sub_heading(doc: DocxDocument, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(paragraph)
    add_mixed_run(paragraph, f"  {text}", east_asia=SONG, size=10.5, bold=True)


def set_cell_text(cell, text: str, east_asia: str = FANGSONG, size: float = 9, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=east_asia, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_title(doc: DocxDocument) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run("数字化转型与企业生产率：一个合成数据工作流示范")
    set_run_font(run, east_asia=SONG, size=16, bold=True)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("经济学研究工作流讲座")
    set_run_font(run, east_asia=SONG, size=10.5)
    mark = author.add_run("*")
    set_run_font(mark, east_asia=SONG, size=9)
    mark.font.superscript = True

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note.paragraph_format.space_after = Pt(6)
    run = note.add_run("* 作者信息：本文为“一站式科研：VS Code 与大模型实操应用”讲座的教学示范文稿。所有数据均为合成数据，结论仅用于工作流展示。")
    set_run_font(run, east_asia=SONG, size=7.5)


def add_abstract(doc: DocxDocument) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = Pt(20)
    head = paragraph.add_run("内容提要：")
    set_run_font(head, east_asia=HEITI, size=12, bold=True)
    body = paragraph.add_run(
        "本文使用合成企业面板数据展示一个从实证估计、理论机制到正式写作的经济学研究工作流。"
        "实证部分采用双重差分设计，估计数字化转型状态与企业对数全要素生产率之间的关系；"
        "理论部分使用 Matlab 构造企业 AI 采用的成本收益优化模型，说明管理能力和采用成本如何影响企业采用 AI 的强度与生产率收益。"
        "所有数据均为模拟生成，本文结果仅用于教学和工作流演示，不能解释为真实企业数字化转型的因果证据。"
    )
    set_run_font(body, east_asia=FANGSONG, size=12)

    keywords = doc.add_paragraph()
    keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    keywords.paragraph_format.line_spacing = Pt(20)
    head = keywords.add_run("关键词：")
    set_run_font(head, east_asia=HEITI, size=12, bold=True)
    body = keywords.add_run("数字化转型；企业生产率；双重差分；事件研究；AI 采用；合成数据")
    set_run_font(body, east_asia=FANGSONG, size=12)


def add_regression_table(doc: DocxDocument) -> None:
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run("表1  DID 主估计结果")
    set_run_font(run, east_asia=SONG, size=9, bold=True)

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    rows = [
        ("Panel A：基准回归：数字化转型与企业生产率", ""),
        ("", "被解释变量"),
        ("变量", "企业对数全要素生产率"),
        ("", "（1）"),
        ("数字化转型", "0.1209***"),
        ("", "（0.0047）"),
        ("资本强度", "0.0121"),
        ("", "（0.0155）"),
        ("出口份额", "-0.2232***"),
        ("", "（0.0153）"),
        ("国有企业", "0.0360***"),
        ("", "（0.0031）"),
        ("企业固定效应", "是"),
        ("年份固定效应", "是"),
        ("观测值", "3,240"),
        ("R²", "0.961"),
    ]

    for row_index, row_data in enumerate(rows):
        row = table.add_row()
        set_cell_text(row.cells[0], row_data[0], east_asia=FANGSONG, size=9, bold=row_index < 4)
        set_cell_text(row.cells[1], row_data[1], east_asia=FANGSONG, size=9, bold=row_index < 4)
        if row_index == 0:
            row.cells[0].merge(row.cells[1])
            set_cell_text(row.cells[0], row_data[0], east_asia=FANGSONG, size=9, bold=True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    note.paragraph_format.first_line_indent = Pt(0)
    note.paragraph_format.line_spacing = Pt(12)
    run = note.add_run("注：括号内为聚类到企业层面的标准误。***、**、* 分别表示在 1%、5% 和 10% 水平上显著。模型控制企业固定效应和年份固定效应。数据为合成数据。")
    set_run_font(run, east_asia=SONG, size=7.5)


def add_figure(doc: DocxDocument, image_name: str, caption_text: str) -> None:
    image_path = OUTPUT / image_name
    if not image_path.exists():
        print(f"Warning: missing figure {image_path}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(5.6))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(caption_text)
    set_run_font(run, east_asia=SONG, size=9)


def add_references(doc: DocxDocument) -> None:
    add_section_heading(doc, "参考文献")
    refs = [
        ("Acemoglu, D. and P. Restrepo, 2019, “Automation and New Tasks: How Technology Displaces and Reinstates Labor,” ", "Journal of Economic Perspectives", ", 33(2): 3–30."),
        ("Agrawal, A., J. S. Gans and A. Goldfarb, 2019, “Artificial Intelligence: The Ambiguous Labor Market Impact of Automating Prediction,” ", "Journal of Economic Perspectives", ", 33(2): 31–50."),
        ("Bertrand, M., E. Duflo and S. Mullainathan, 2004, “How Much Should We Trust Differences-in-Differences Estimates?” ", "The Quarterly Journal of Economics", ", 119(1): 249–275."),
        ("Bloom, N., R. Sadun and J. Van Reenen, 2012, “Americans Do IT Better: US Multinationals and the Productivity Miracle,” ", "American Economic Review", ", 102(1): 167–201."),
        ("Callaway, B. and P. H. C. Sant’Anna, 2021, “Difference-in-Differences with Multiple Time Periods,” ", "Journal of Econometrics", ", 225(2): 200–230."),
        ("Goodman-Bacon, A., 2021, “Difference-in-Differences with Variation in Treatment Timing,” ", "Journal of Econometrics", ", 225(2): 254–277."),
        ("Sun, L. and S. Abraham, 2021, “Estimating Dynamic Treatment Effects in Event Studies with Heterogeneous Treatment Effects,” ", "Journal of Econometrics", ", 225(2): 175–199."),
    ]
    for before, journal, after in refs:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(10.5)
        paragraph.paragraph_format.line_spacing = Pt(12)
        run = paragraph.add_run("  " + before)
        set_run_font(run, east_asia=SONG, size=7.5)
        journal_run = paragraph.add_run(journal)
        set_run_font(journal_run, east_asia=SONG, size=7.5, italic=True)
        run = paragraph.add_run(after)
        set_run_font(run, east_asia=SONG, size=7.5)


def build_document() -> None:
    doc = create_document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    styles = doc.styles
    normal = cast(Any, styles["Normal"])
    normal.font.name = TIMES
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), SONG)
    normal.font.size = Pt(10.5)

    add_title(doc)
    add_abstract(doc)

    add_section_heading(doc, "一、引言")
    add_body_paragraph(doc, "数字化转型和人工智能应用正在改变企业组织生产、处理信息和配置资源的方式。对经济学研究者而言，一个自然问题是：采用数字技术的企业是否在采用后表现出更高的生产率？更进一步，企业是否会因为采用成本和管理能力差异而做出不同的技术采用决策？本文围绕这些问题构造一个可复现的教学示范。示范的核心不是提供真实世界的估计结论，而是展示如何把 Stata 中的双重差分估计、Matlab 中的机制模型、Jupyter notebook 中的结果汇总和 LaTeX 中的正式写作连接起来。")
    add_body_paragraph(doc, "本文的研究对象是一个合成企业面板。处理组企业从 2020 年开始进入数字化转型状态，对照组企业没有进入该状态。实证模型控制企业固定效应和年份固定效应，并将标准误聚类到企业层面。主要估计结果显示，数字化转型状态与企业对数生产率之间的相关 DID 系数约为 0.1209，企业层面聚类标准误约为 0.0047。由于数据是人为构造的，这个数值只用于检查工作流是否能把分析结果传递到论文写作环节。")

    add_section_heading(doc, "二、文献综述")
    add_body_paragraph(doc, "本文的示范连接了两类文献。第一类文献讨论信息技术、自动化和人工智能对生产率与劳动市场的影响。Bloom et al.（2012）研究跨国企业的信息技术使用与生产率差异，强调管理和组织因素在 IT 收益中的作用。Acemoglu & Restrepo（2019）讨论自动化对任务结构的替代与创造，Agrawal et al.（2019）则从预测成本下降的角度理解人工智能的经济影响。本文的 Matlab 模型沿用这一思路，将 AI 采用看作企业在生产率收益和实施成本之间的权衡。")
    add_body_paragraph(doc, "第二类文献关注双重差分和事件研究方法。Bertrand et al.（2004）指出 DID 推断中序列相关和标准误处理的重要性，因此本文将标准误聚类到企业层面。对于多期处理和动态效应设定，Callaway & Sant’Anna（2021）、Sun & Abraham（2021）和 Goodman-Bacon（2021）提醒研究者注意处理时点异质性和事件研究解释。本文的事件研究图只是合成数据中的识别诊断示范，不构成真实研究中的识别证明。")

    add_section_heading(doc, "三、数据与变量")
    add_body_paragraph_with_math(doc, "合成数据覆盖 2016 至 2024 年的企业面板。核心结果变量为企业对数全要素生产率 $\\log(TFP_{it})$。核心解释变量 $Digital_{it}$ 表示企业是否已经进入数字化转型状态。控制变量包括资本强度、出口份额和所有制变量。处理组企业在 2020 年后进入数字化转型状态，因此该变量同时包含处理组身份和处理后时期信息。")
    add_body_paragraph(doc, "再次强调，样本并不来自真实企业数据库。数据生成过程被设计为服务于现场演示：一方面，它要足够接近经济学论文的分析流程；另一方面，它必须可被 Python、Stata、Matlab、Jupyter 和 LaTeX 共同读取和复现。")

    add_section_heading(doc, "四、实证策略")
    add_body_paragraph_with_math(doc, "基准模型为双重差分设定：")
    add_display_equation(doc, "\\log(TFP_{it}) = \\alpha_i + \\lambda_t + \\tau Digital_{it} + X_{it}'\\beta + \\varepsilon_{it},")
    add_body_paragraph_with_math(doc, "其中 $\\alpha_i$ 表示企业固定效应，$\\lambda_t$ 表示年份固定效应，$X_{it}$ 包括资本强度、出口份额和所有制控制变量。标准误聚类到企业层面，以反映同一企业内误差项可能存在相关性。")
    add_body_paragraph(doc, "为了展示动态效应，本文还构造相对采用年份的事件研究变量，并以采用前一年作为基准组。事件研究用于检查处理前系数是否接近零，以及处理后系数是否按模拟设定逐步上升。它是识别诊断，而不是识别假设本身的证明。")

    add_section_heading(doc, "五、主要结果")
    add_body_paragraph(doc, "表1列（1）报告主要结果。数字化转型变量的估计系数为 0.1209，聚类标准误为 0.0047，说明在合成数据中，处理组企业进入数字化转型状态后，对数生产率相对于对照组出现上升。这个结果与数据生成过程一致，但不能被解释为真实企业数字化转型的平均处理效应。")
    add_regression_table(doc)

    add_section_heading(doc, "六、动态效应与识别检查")
    add_body_paragraph(doc, "图1展示处理组与对照组在处理前后的平均趋势。图2展示以处理前一年为基准的事件研究系数和 95% 置信区间。处理前相对年份 -4、-3 和 -2 的系数分别约为 0.0034、0.0041 和 0.0108，均接近零；处理后系数从 0.0573 上升到 0.1974。这一模式符合模拟数据的设计：处理前差异较小，处理后效应逐步增强。")
    add_figure(doc, "digital_parallel_trends.png", "图1  合成企业面板中的平行趋势诊断")
    add_figure(doc, "stata_event_study_ci.png", "图2  事件研究估计与 95% 置信区间")

    add_section_heading(doc, "七、安慰剂检验")
    add_body_paragraph(doc, "安慰剂检验只使用处理前样本，并假设处理组企业在 2018 年已经受到虚假处理。估计结果显示，placebo_digital_2018 的系数约为 0.0015，标准误约为 0.0074，p 值约为 0.835。这说明在合成数据中，处理前没有明显的虚假政策效应。该检验支持工作流中的诊断逻辑，但仍不能替代真实研究中对识别假设的论证。")

    add_section_heading(doc, "八、理论机制：AI 采用成本收益模型")
    add_body_paragraph_with_math(doc, "Matlab 部分将数字化转型解释为企业 AI 采用决策。企业 $i$ 的管理能力为 $m_i$，选择 AI 使用强度 $x_i \\in [0,1]$。采用 AI 的净收益为")
    add_display_equation(doc, "V_i(x_i) = (\\phi + \\lambda m_i)x_i - \\frac{1}{2}\\psi x_i^2 - c_i.")
    add_body_paragraph_with_math(doc, "企业先使用有界优化求解最优强度 $x_i^*$，再根据 $V_i(x_i^*)$ 是否大于零决定是否采用 AI。模型用 Stata DID 估计目标校准参数 $\\phi$。校准结果显示，$\\phi$ 约为 0.8659，模型平均收益约为 0.1209，AI 采用率约为 0.8197。")
    add_body_paragraph(doc, "图3展示不同管理能力企业的模型预测生产率收益。管理能力较高的企业因为采用成本相对更低、AI 使用回报更高，更容易采用 AI 并获得更高收益。该模型只是机制展示，不是对真实企业行为的结构估计。")
    add_figure(doc, "matlab_theory_model.png", "图3  AI 采用优化模型中的生产率收益")

    add_section_heading(doc, "九、讨论与局限")
    add_body_paragraph(doc, "本文最大的局限是数据完全由模拟生成。因此，所有估计值只说明工作流从数据生成、Stata 回归、Matlab 机制模拟到论文写作之间是连贯的，不能说明真实企业采用数字技术后一定会提高生产率。第二，事件研究和平行趋势图是诊断工具，不是识别假设的形式证明。第三，理论模型被有意设计得很简洁，只展示成本收益权衡，并未估计真实企业的技术采用结构参数。")

    add_section_heading(doc, "十、结论")
    add_body_paragraph(doc, "本文展示了一个完整的经济学研究写作工作流。合成数据中的 DID 估计、事件研究、安慰剂检验和 Matlab 机制模型共同构成一个可复现示范。更重要的是，该示范说明 AI 工具可以辅助研究者把可执行分析转化为正式论文文本，但研究者仍必须核查文献、验证代码、检查编译结果，并明确区分演示性结果和真实经验证据。")

    add_section_heading(doc, "十一、AI 写作与复现审计附录")
    add_body_paragraph(doc, "本文写作使用了 AI 辅助工具读取 notebook、输出文件和 LaTeX 初稿，并生成中英文论文草稿。研究者保留了合成数据限制、估计对象、聚类标准误、识别诊断和文献核查要求。所有引用均来自 DOI 或 Crossref 可核查条目；未能核查的文献没有进入参考文献列表。LaTeX 源文件通过 XeLaTeX 编译验证，审计记录见项目根目录下的 audit-log.md。")

    add_references(doc)
    doc.save(str(DOCX_PATH))


if __name__ == "__main__":
    build_document()
    print(DOCX_PATH)